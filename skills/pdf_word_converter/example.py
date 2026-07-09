"""
PDF-Word Converter 使用示例
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from  pdf_word_converter import PdfWordConverter


def example_check_dependencies():
    """检查依赖"""
    print("=== 检查依赖 ===\n")
    converter = PdfWordConverter()

    print(f"PDF转Word: {'可用' if converter.is_pdf_to_word_available() else '不可用'}")
    print(f"Word转PDF: {'可用' if converter.is_word_to_pdf_available() else '不可用'}")

    if not converter.is_pdf_to_word_available():
        print("\n请安装: pip install pdf2docx")
    if not converter.is_word_to_pdf_available():
        print("\n请安装: pip install docx2pdf")


def example_basic_usage():
    """基础使用示例"""
    print("\n=== 基础使用 ===\n")

    converter = PdfWordConverter()

    # 创建一个测试用的docx文件（如果有python-docx）
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("测试文档", 0)
        doc.add_paragraph("这是一个用于测试PDF-Word转换的示例文档。")
        doc.add_paragraph("你好，世界！")
        test_docx = "test_document.docx"
        doc.save(test_docx)
        print(f"已创建测试文档: {test_docx}")

        # 测试Word转PDF（如果可用）
        if converter.is_word_to_pdf_available():
            try:
                test_pdf = converter.word_to_pdf(test_docx)
                print(f"Word转PDF成功: {test_pdf}")

                # 再转回Word
                if converter.is_pdf_to_word_available():
                    test_docx2 = converter.pdf_to_word(test_pdf, "test_converted_back.docx")
                    print(f"PDF转回Word成功: {test_docx2}")
            except Exception as e:
                print(f"转换测试遇到问题: {e}")
                print("（这可能是因为缺少Microsoft Word应用程序）")

    except ImportError:
        print("python-docx 未安装，跳过创建测试文档")
        print("你可以使用现有的Word或PDF文件进行测试")


def example_batch_conversion():
    """批量转换示例"""
    print("\n=== 批量转换 ===\n")

    converter = PdfWordConverter()

    print("批量转换文件夹使用方法:")
    print("  converter.convert_folder('input_dir', 'output_dir', mode='pdf2word')")
    print("  mode 可选: 'pdf2word' 或 'word2pdf'")


if __name__ == "__main__":
    example_check_dependencies()

    print("\n" + "="*50 + "\n")

    example_basic_usage()

    print("\n" + "="*50 + "\n")

    example_batch_conversion()

    print("\n提示:")
    print("- Word转PDF功能需要安装Microsoft Word (Windows) 或 LibreOffice")
    print("- 如果没有Word，你可以只使用PDF转Word功能")
